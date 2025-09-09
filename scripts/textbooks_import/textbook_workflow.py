#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
教材课文处理工作流

此脚本整合了教材处理的完整流程：
1. Word文档模块分割
2. 课文结构化提取
3. 背诵时间估算
4. 生成最终的SQL插入语句

使用示例:
    python textbook_workflow.py --input "path/to/textbook.doc" --textbook_id 1
"""

import asyncio
import argparse
import json
import os
import sys
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging
from datetime import datetime
import re
from docx import Document

# 添加项目根目录到Python路径
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.append(str(PROJECT_ROOT))

from app.services.glm4_service import GLM4Service
from app.schemas.glm4 import GLM4Request

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('textbook_workflow.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


def estimate_memorization_time_simple(word_count: int, difficulty_level: int = 2) -> int:
    """
    简化版的英文课文背诵时间估算算法
    
    参数:
    word_count (int): 课文的单词数量
    difficulty_level (int): 课文难度级别 (1-简单, 2-中等, 3-困难)
    
    返回:
    int: 推荐的背诵时长（整数分钟，不足1分钟按1分钟计）
    """
    
    # 基础背诵速度（单词/分钟） - 基于中等水平学生
    base_speed = 3.0
    
    # 难度系数
    difficulty_factors = {1: 0.7, 2: 1.0, 3: 1.5}
    
    # 计算基础记忆时间（分钟）
    base_time = word_count / base_speed
    
    # 应用难度系数
    memorization_time = base_time * difficulty_factors.get(difficulty_level, 1.0)
    
    # 考虑理解内容带来的效率提升（默认学生理解内容）
    comprehension_bonus = 0.7  # 理解内容可减少30%的时间
    
    # 最终推荐时间
    recommended_time = memorization_time * comprehension_bonus
    
    # 向上取整到整数分钟，最少1分钟
    return max(1, int(recommended_time + 0.5))


def count_words(text: str) -> int:
    """
    计算英文文本的单词数量
    
    Args:
        text: 英文文本
        
    Returns:
        int: 单词数量
    """
    if not text:
        return 0
    
    # 使用正则表达式匹配英文单词
    words = re.findall(r'\b[A-Za-z]+\b', text)
    return len(words)


def convert_doc_to_docx(doc_file_path: str) -> str:
    """
    将.doc文件转换为.docx格式
    
    Args:
        doc_file_path (str): .doc文件路径
        
    Returns:
        str: 转换后的.docx文件路径
        
    Raises:
        Exception: 转换失败
    """
    import subprocess
    
    docx_file_path = doc_file_path.rsplit('.', 1)[0] + '.docx'
    
    # 尝试使用LibreOffice进行转换
    try:
        logger.info("尝试使用LibreOffice转换.doc文件...")
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'docx', 
            '--outdir', os.path.dirname(doc_file_path), doc_file_path
        ], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists(docx_file_path):
            logger.info(f"✓ 已使用LibreOffice将.doc文件转换为.docx格式: {docx_file_path}")
            return docx_file_path
            
    except FileNotFoundError:
        pass  # LibreOffice不可用，继续尝试其他方法
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice转换超时")
    except Exception as e:
        logger.warning(f"LibreOffice转换失败: {str(e)}")
    
    # 尝试使用textutil (macOS自带工具)
    try:
        logger.info("尝试使用textutil转换.doc文件...")
        # 先转换为rtf格式，再用pandoc转换为docx
        rtf_file_path = doc_file_path.rsplit('.', 1)[0] + '.rtf'
        
        # .doc -> .rtf
        result1 = subprocess.run([
            'textutil', '-convert', 'rtf', doc_file_path, '-output', rtf_file_path
        ], capture_output=True, text=True)
        
        if result1.returncode == 0 and os.path.exists(rtf_file_path):
            # .rtf -> .docx
            result2 = subprocess.run([
                'pandoc', rtf_file_path, '-o', docx_file_path
            ], capture_output=True, text=True)
            
            # 清理临时rtf文件
            if os.path.exists(rtf_file_path):
                os.remove(rtf_file_path)
                
            if result2.returncode == 0 and os.path.exists(docx_file_path):
                logger.info(f"✓ 已使用textutil+pandoc将.doc文件转换为.docx格式: {docx_file_path}")
                return docx_file_path
                
    except FileNotFoundError:
        pass  # textutil不可用（非macOS系统）
    except Exception as e:
        logger.warning(f"textutil转换失败: {str(e)}")
    
    # 所有自动转换方法都失败了，给出详细的手动转换指导
    raise Exception(
        f"❌ 无法自动转换.doc文件。Pandoc不支持.doc格式（只支持.docx）。\n\n"
        f"请选择以下解决方案之一：\n\n"
        f"🔧 解决方案1：使用Microsoft Word\n"
        f"   1. 用Word打开文件：{doc_file_path}\n"
        f"   2. 选择'文件' -> '另存为'\n"
        f"   3. 选择格式为'Word文档(*.docx)'\n"
        f"   4. 保存后重新运行脚本\n\n"
        f"🔧 解决方案2：安装LibreOffice (免费)\n"
        f"   1. 访问: https://www.libreoffice.org/download/\n"
        f"   2. 安装后重新运行脚本，将自动转换\n\n"
        f"🔧 解决方案3：使用在线转换工具\n"
        f"   - https://convertio.co/doc-docx/\n"
        f"   - https://www.zamzar.com/convert/doc-to-docx/\n\n"
        f"💡 提示：转换后的.docx文件可以直接被脚本处理"
    )


def read_word_document(file_path: str) -> str:
    """
    读取Word文档内容，自动处理.doc和.docx格式
    
    Args:
        file_path (str): Word文档文件路径
        
    Returns:
        str: 文档的纯文本内容
        
    Raises:
        FileNotFoundError: 文件不存在
        Exception: 文档读取失败
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    original_path = file_path
    file_extension = file_path.lower().split('.')[-1]
    
    if file_extension not in ['doc', 'docx']:
        raise ValueError("仅支持.doc和.docx格式的Word文档")
    
    # 如果是.doc格式，尝试转换为.docx
    if file_extension == 'doc':
        try:
            file_path = convert_doc_to_docx(file_path)
        except Exception as e:
            logger.error(f"转换.doc文件失败: {str(e)}")
            raise
    
    try:
        doc = Document(file_path)
        full_text = []
        
        # 读取所有段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 忽略空段落
                full_text.append(paragraph.text)
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        # 如果我们转换了.doc文件，清理临时文件
        if original_path != file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"已清理临时文件: {file_path}")
            except:
                pass
        
        return '\n'.join(full_text)
        
    except Exception as e:
        raise Exception(f"读取文档失败: {str(e)}")


def split_by_module(content: str, module_keyword: str = "Module") -> List[str]:
    """
    按照Module关键词分割文档内容
    
    Args:
        content (str): 文档内容
        module_keyword (str): 分割关键词，默认为"Module"
        
    Returns:
        List[str]: 分割后的内容列表，每个元素代表一个Module的内容
    """
    if not content.strip():
        return []
    
    # 使用正则表达式查找Module关键词的位置
    # 匹配模式：行首的Module（可能前面有数字、空格等）
    pattern = rf'^.*?{re.escape(module_keyword)}\s*\d*.*?$'
    matches = list(re.finditer(pattern, content, re.MULTILINE | re.IGNORECASE))
    
    if not matches:
        # 如果没有找到Module关键词，返回整个文档作为一个部分
        logger.warning(f"未找到'{module_keyword}'关键词，返回完整文档")
        return [content]
    
    modules = []
    
    for i, match in enumerate(matches):
        start_pos = match.start()
        
        # 确定当前Module的结束位置
        if i + 1 < len(matches):
            end_pos = matches[i + 1].start()
        else:
            end_pos = len(content)
        
        # 提取当前Module的内容
        module_content = content[start_pos:end_pos].strip()
        if module_content:
            modules.append(module_content)
    
    return modules


class TextbookWorkflow:
    """教材课文处理工作流"""
    
    def __init__(self, textbook_path: str, textbook_id: int):
        """
        初始化工作流
        
        Args:
            textbook_path: 教材文档路径
            textbook_id: 教材ID
        """
        self.textbook_path = textbook_path
        self.textbook_id = textbook_id
        self.glm4_service = GLM4Service()
        self.temp_dir = None
        self.results_dir = Path("scripts/textbooks_import/results")
        self.results_dir.mkdir(exist_ok=True)
        
    def __enter__(self):
        """进入上下文管理器"""
        self.temp_dir = tempfile.mkdtemp(prefix="textbook_workflow_")
        logger.info(f"创建临时工作目录: {self.temp_dir}")
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        """退出上下文管理器"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
            logger.info(f"清理临时工作目录: {self.temp_dir}")
    
    async def process_textbook(self) -> str:
        """
        处理教材的完整流程
        
        Returns:
            str: 生成的SQL文件路径
        """
        logger.info(f"开始处理教材: {self.textbook_path}")
        
        # 步骤1：读取和分割文档
        logger.info("步骤1: 读取和分割Word文档")
        content = read_word_document(self.textbook_path)
        logger.info(f"文档读取成功，总字符数: {len(content)}")
        
        modules = split_by_module(content)
        logger.info(f"文档分割完成，共找到 {len(modules)} 个Module")
        
        # 步骤2：提取课文信息
        logger.info("步骤2: 使用GLM-4提取课文结构")
        all_lessons = []
        
        for i, module_content in enumerate(modules, 1):
            module_name = f"module_{i:02d}"
            logger.info(f"正在处理 {module_name}...")
            
            lessons = await self.extract_lessons_from_module(module_content, module_name)
            if lessons:
                all_lessons.extend(lessons)
                logger.info(f"从 {module_name} 提取到 {len(lessons)} 个课文")
            else:
                logger.warning(f"{module_name} 未提取到课文内容")
            
            # 添加延迟以避免API限制
            await asyncio.sleep(2)
        
        logger.info(f"总计提取到 {len(all_lessons)} 个课文")
        
        # 步骤3：计算背诵时间并生成SQL
        logger.info("步骤3: 计算背诵时间并生成SQL语句")
        sql_file_path = self.generate_final_sql(all_lessons)
        
        logger.info("教材处理完成！")
        return sql_file_path
    
    async def extract_lessons_from_module(self, module_content: str, module_name: str, max_retries: int = 3) -> List[Dict[str, Any]]:
        """
        从模块内容中提取课文信息（支持重试）
        
        Args:
            module_content: 模块内容
            module_name: 模块名称
            max_retries: 最大重试次数，默认3次
            
        Returns:
            List[Dict]: 提取的课文信息列表
        """
        
        for attempt in range(max_retries + 1):  # 总共尝试 max_retries + 1 次
            try:
                if attempt > 0:
                    logger.info(f"正在重试处理 {module_name}... (第{attempt}/{max_retries}次重试)")
                    # 重试时稍微增加延迟
                    await asyncio.sleep(3)
                
                prompt = self.get_extraction_prompt(module_content)
                request = GLM4Request(prompt=prompt, temperature=0.3)
                
                response = await self.glm4_service.basic_call(request)
                
                # 提取响应内容
                if response.choices and len(response.choices) > 0:
                    message = response.choices[0].message
                    content = message.get('content', '') if isinstance(message, dict) else str(message)
                    
                    # 解析JSON
                    try:
                        result = json.loads(content)
                        lessons = result.get('lessons', [])
                        if lessons:  # 只有成功提取到课文才算成功
                            if attempt > 0:
                                logger.info(f"✓ {module_name} 重试成功，提取到{len(lessons)}个课文")
                            return lessons
                        else:
                            logger.warning(f"GLM-4返回空课文列表: {module_name}")
                            if attempt < max_retries:
                                continue  # 继续重试
                            else:
                                return []
                                
                    except json.JSONDecodeError as e:
                        logger.error(f"JSON解析失败 {module_name} (尝试{attempt + 1}/{max_retries + 1}): {str(e)}")
                        logger.error(f"响应内容: {content}")
                        if attempt < max_retries:
                            continue  # 继续重试
                        else:
                            return []
                else:
                    logger.error(f"GLM-4响应为空: {module_name} (尝试{attempt + 1}/{max_retries + 1})")
                    if attempt < max_retries:
                        continue  # 继续重试
                    else:
                        return []
                        
            except Exception as e:
                logger.error(f"提取课文信息失败 {module_name} (尝试{attempt + 1}/{max_retries + 1}): {str(e)}")
                if attempt < max_retries:
                    continue  # 继续重试
                else:
                    return []
        
        return []
    
    def get_extraction_prompt(self, module_content: str) -> str:
        """
        构建GLM-4提取提示词
        
        Args:
            module_content: 模块内容
            
        Returns:
            str: GLM-4提示词
        """
        prompt = f"""
请分析以下英语教材课文模块内容，提取并校正其结构化信息。

原始内容：
{module_content}

请按照以下JSON格式输出每个课文单元的信息，注意校正英文和翻译中的错误：

{{
  "lessons": [
    {{
      "unit_number": "Module X Unit Y",
      "unit_title": "单元英文标题",
      "lesson_title": "课文类型标题（如：1. Listen and chant）",
      "title": "课文标题（如：Listen and chant）",
      "content": "英文课文内容（校正后的，按语义分段）",
      "translation": "中文翻译（校正后的，按语义分段）",
      "difficulty_level": 2
    }}
  ]
}}

提取要求：
1. 准确识别每个Unit和lesson的边界
2. 校正英文语法、拼写错误
3. 校正中文翻译错误，确保通顺准确
4. content字段只包含英文原文，translation字段只包含中文翻译
5. **语义分段要求**：
   - 对于content和translation字段，请按照语义进行合理分段
   - 使用换行符（\\n）分隔不同的语义段落
   - 分段原则：
     * 对话内容：每个说话者的话为一段
     * 故事/叙述：按照意思相关的句群分段
     * 歌谣/诗歌：按照韵律或意思分行
     * 描述性文字：按照主题或场景分段
   - 确保分段后的内容逻辑清晰，便于阅读和背诵
6. difficulty_level评估课文难度：1-简单（单词简单、句式简单），2-中等（适中难度），3-困难（词汇复杂、语法复杂）
7. 保持课文的完整性和准确性
8. 确保输出的是标准的JSON格式

分段示例：
- 对话形式："Hello, Amy.\\nHello, Sam.\\nHow are you?\\nI'm fine, thank you."
- 叙述形式："This is my family.\\nMy father is a doctor.\\nMy mother is a teacher.\\nI love my family."
- 歌谣形式："Ten little fingers,\\nTen little toes,\\nTwo little ears,\\nAnd one little nose."

请直接输出JSON，不要包含其他说明文字。
"""
        return prompt
    
    def generate_final_sql(self, lessons: List[Dict[str, Any]]) -> str:
        """
        生成包含recommend_minutes字段的最终SQL语句
        
        Args:
            lessons: 课文信息列表
            
        Returns:
            str: SQL文件路径
        """
        if not lessons:
            logger.warning("没有课文数据，无法生成SQL")
            return ""
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        sql_file = self.results_dir / f"textbooks_lessons_insert_{timestamp}.sql"
        
        # 生成SQL内容
        sql_parts = []
        sql_parts.append("-- 教材课文数据插入语句")
        sql_parts.append(f"-- 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        sql_parts.append(f"-- 教材ID: {self.textbook_id}")
        sql_parts.append(f"-- 总课文数: {len(lessons)}")
        sql_parts.append("")
        sql_parts.append("INSERT INTO public.textbooks_lessons (")
        sql_parts.append("  textbook_id, unit_number, unit_title, lesson_title,")
        sql_parts.append("  lesson_order, title, content, translation,")
        sql_parts.append("  recommend_minutes, author, source, created_at, updated_at")
        sql_parts.append(") VALUES")
        
        value_parts = []
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # lesson_order 从1开始递增，表示课文在整本教材中的序号
        for lesson_order, lesson in enumerate(lessons, 1):
            # 计算背诵时间
            content = lesson.get('content', '')
            word_count = count_words(content)
            difficulty_level = lesson.get('difficulty_level', 2)  # 使用模型输出的难度等级，默认为2
            recommend_minutes = estimate_memorization_time_simple(word_count, difficulty_level=difficulty_level)
            
            # 处理字符串中的单引号，防止SQL注入
            def escape_sql_string(s: str) -> str:
                if s is None:
                    return 'NULL'
                return "'" + str(s).replace("'", "''") + "'"
            
            value_part = f"""(
  {self.textbook_id},
  {escape_sql_string(lesson.get('unit_number', ''))},
  {escape_sql_string(lesson.get('unit_title', ''))},
  {escape_sql_string(lesson.get('lesson_title', ''))},
  {lesson_order},
  {escape_sql_string(lesson.get('title', ''))},
  {escape_sql_string(content)},
  {escape_sql_string(lesson.get('translation', ''))},
  {recommend_minutes},
  '外研社',
  '教材导入',
  '{current_time}',
  '{current_time}'
)"""
            value_parts.append(value_part)
        
        sql_parts.append(",\n".join(value_parts))
        sql_parts.append(";")
        
        # 保存SQL文件
        with open(sql_file, 'w', encoding='utf-8') as f:
            f.write("\n".join(sql_parts))
        
        logger.info(f"SQL语句已保存: {sql_file}")
        
        # 生成处理摘要
        summary = []
        summary.append(f"处理完成摘要:")
        summary.append(f"- 教材路径: {self.textbook_path}")
        summary.append(f"- 教材ID: {self.textbook_id}")
        summary.append(f"- 总课文数: {len(lessons)}")
        summary.append(f"- SQL文件: {sql_file}")
        summary.append("")
        summary.append("课文详情:")
        
        for i, lesson in enumerate(lessons, 1):
            content = lesson.get('content', '')
            word_count = count_words(content)
            difficulty_level = lesson.get('difficulty_level', 2)  # 使用模型输出的难度等级，默认为2
            recommend_minutes = estimate_memorization_time_simple(word_count, difficulty_level=difficulty_level)
            
            difficulty_text = {1: '简单', 2: '中等', 3: '困难'}.get(difficulty_level, '中等')
            summary.append(f"  {i}. {lesson.get('title', 'N/A')} "
                         f"({word_count}词, 难度:{difficulty_text}, {recommend_minutes}分钟)")
        
        summary_text = "\n".join(summary)
        logger.info(f"\n{summary_text}")
        
        # 保存处理摘要
        summary_file = self.results_dir / f"processing_summary_{timestamp}.txt"
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write(summary_text)
        
        return str(sql_file)


async def main(file_path: str, textbook_id: int):
    """主程序"""
    
    # 验证输入文件
    if not os.path.exists(file_path):
        logger.error(f"输入文件不存在: {file_path}")
        sys.exit(1)
    
    try:
        with TextbookWorkflow(file_path, textbook_id) as workflow:
            sql_file_path = await workflow.process_textbook()
            
            print(f"\n🎉 处理完成！")
            print(f"📄 输入文件: {file_path}")
            print(f"🆔 教材ID: {textbook_id}")
            print(f"📝 生成的SQL文件: {sql_file_path}")
            print(f"\n💡 下一步：将SQL文件导入到数据库中")
            
    except Exception as e:
        logger.error(f"处理失败: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    filepath = "/Users/ethan/PycharmProjects/polly-memo-fastapi/scripts/textbooks_import/5下 外研一起 课文.doc"
    textbook_id = 2
    asyncio.run(main(filepath, textbook_id)) 